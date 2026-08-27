"""中文评估报告生成（Markdown）：概况、设计降雨、结果、受影响区域分析、声明。"""

from __future__ import annotations

from datetime import datetime


def _sample(flood: list, bbox: list, lon: float, lat: float) -> float:
    """积水场最近邻采样（度 → 网格索引）。"""
    h, w = len(flood), len(flood[0])
    west, south, east, north = bbox
    c = (lon - west) / (east - west) * (w - 1)
    r = (north - lat) / (north - south) * (h - 1)
    if not (0 <= r < h and 0 <= c < w):
        return 0.0
    return flood[int(r)][int(c)]


def analyze_impact(viz: dict) -> dict:
    """OSM 地名/主干道 × 积水场采样 → 受影响社区与道路（按深度降序）。"""
    fl, bbox = viz["flood"], viz["bbox"]
    osm = viz.get("osm") or {}
    communities = []
    for lon, lat, name in osm.get("places", []):
        d = _sample(fl, bbox, lon, lat)
        if d >= 20:
            communities.append({"name": name, "depth_mm": round(d)})
    communities.sort(key=lambda x: -x["depth_mm"])
    roads = []
    for line in osm.get("road_major", []):
        mx = max(_sample(fl, bbox, lo, la) for lo, la in line)
        if mx >= 30:
            roads.append({"name": None, "depth_mm": round(mx)})
    roads.sort(key=lambda x: -x["depth_mm"])
    return {"communities": communities[:10], "roads": roads[:8]}


def generate_markdown(r: dict) -> str:
    viz = r["viz"]
    impact = analyze_impact(viz)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    comm = impact["communities"]
    comm_md = "\n".join(
        f"{i+1}. **{c['name']}**：积水约 {c['depth_mm']} mm" for i, c in enumerate(comm)
    ) or "未识别到积水深度 ≥20mm 的社区"
    top5 = "\n".join(
        f"| {i+1} | {t['node']} | {t['depth_mm']} mm |" for i, t in enumerate(r["top5"])
    )
    imp_lo, imp_hi = r["imperv_range_pct"]
    return f"""# 城市内涝快速评估报告

> 生成时间：{now} · swmm-copilot（内涝快评）

## 一、评估概况

| 项目 | 内容 |
|---|---|
| 评估片区 | {r['aoi']}（{r['city']}） |
| 暴雨公式分区 | {r['formula_zone']}（依据当地暴雨强度公式） |
| 设计重现期 | {r['return_period_yr']} 年一遇 |
| 评估方法 | DEM 填洼 + D8 汇流分析 + SWMM 网格概化模拟（管网概化生成） |
| 数据来源 | Copernicus GLO-30 DEM · ESA WorldCover 土地覆盖 · OSM 路网水系地名 |

## 二、设计降雨

- 设计暴雨历时 120 分钟（芝加哥雨型，雨峰系数 0.4）
- **设计总雨量：{r['rain_total_mm']} mm**
- 片区不透水率范围：{imp_lo:.0f}% ~ {imp_hi:.0f}%（WorldCover 建成区估算）

## 三、模拟结果

- 管网节点共 {r['total_nodes']} 个，其中 **{r['flood_nodes']} 个发生溢流**
- **最大折算积水深度：{r['max_depth_mm']:.0f} mm**

积水最深点 Top 5：

| 排名 | 节点 | 积水深度 |
|---|---|---|
{top5}

## 四、受影响区域分析

依据 OSM 地名与主干路网叠加积水场采样：

{comm_md}

> 注：积水范围为由节点溢流深度插值的示意性淹没场，受影响判定阈值：社区 ≥20mm、道路 ≥30mm。

## 五、结论与建议

在 {r['return_period_yr']} 年一遇设计降雨下，{r['aoi']} 片区溢流节点占比
{r['flood_nodes']}/{r['total_nodes']}，建议优先核查上表积水最深区域的排水干管
能力与下游顶托情况，并结合低洼社区（{comm[0]['name'] if comm else '—'} 等）
开展积水点整治与应急排涝预案复核。

## 附：成果文件

- 积水风险图：`flood_map.png`（同目录）
- SWMM 模型：`model.inp`（同目录，可在 SWMM/pyswmm 中复算）

---

**声明**：本报告为**规划级快速评估**，{r['note']}。由 swmm-copilot 自动生成，正式工程请以实测管网与详细设计为准。
"""


def generate_docx(r: dict, path) -> None:
    """Word 版评估报告（内容与 Markdown 版一致）。依赖 python-docx。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    viz = r["viz"]
    impact = analyze_impact(viz)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_heading("城市内涝快速评估报告", 0)
    doc.add_paragraph(f"评估片区：{r['aoi']}（{r['city']}） · {r['return_period_yr']} 年一遇 · swmm-copilot")

    doc.add_heading("一、评估概况", 1)
    t = doc.add_table(rows=5, cols=2)
    t.style = "Light Grid Accent 1"
    rows = [("暴雨公式分区", f"{r['formula_zone']}（依据当地暴雨强度公式）"),
            ("设计重现期", f"{r['return_period_yr']} 年一遇"),
            ("评估方法", "DEM 填洼 + D8 汇流分析 + SWMM 网格概化模拟"),
            ("数据来源", "Copernicus GLO-30 · ESA WorldCover · OSM")]
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text, t.cell(i, 1).text = k, v

    doc.add_heading("二、设计降雨", 1)
    doc.add_paragraph(f"设计暴雨历时 120 分钟（芝加哥雨型，雨峰系数 0.4），设计总雨量 {r['rain_total_mm']} mm；"
                      f"片区不透水率 {r['imperv_range_pct'][0]:.0f}% ~ {r['imperv_range_pct'][1]:.0f}%。")

    doc.add_heading("三、模拟结果", 1)
    doc.add_paragraph(f"管网节点共 {r['total_nodes']} 个，其中 {r['flood_nodes']} 个发生溢流；"
                      f"最大折算积水深度 {r['max_depth_mm']:.0f} mm。")
    t2 = doc.add_table(rows=6, cols=3)
    t2.style = "Light Grid Accent 1"
    for j, h in enumerate(("排名", "节点", "积水深度")):
        t2.cell(0, j).text = h
    for i, item in enumerate(r["top5"], 1):
        t2.cell(i, 0).text, t2.cell(i, 1).text, t2.cell(i, 2).text = \
            str(i), item["node"], f"{item['depth_mm']} mm"

    doc.add_heading("四、受影响区域分析", 1)
    for i, c in enumerate(impact["communities"], 1):
        doc.add_paragraph(f"{i}. {c['name']}：积水约 {c['depth_mm']} mm", style="List Number")
    doc.add_paragraph("注：积水范围为节点溢流深度插值的示意性淹没场（社区阈值 ≥20mm、道路 ≥30mm）。")

    doc.add_heading("五、结论与建议", 1)
    top_comm = impact["communities"][0]["name"] if impact["communities"] else "—"
    doc.add_paragraph(
        f"在 {r['return_period_yr']} 年一遇设计降雨下，{r['aoi']} 片区溢流节点占比 "
        f"{r['flood_nodes']}/{r['total_nodes']}。建议优先核查积水最深区域的排水干管能力与下游顶托情况，"
        f"并结合低洼社区（{top_comm} 等）开展积水点整治与应急排涝预案复核。")

    doc.add_paragraph(f"声明：本报告为规划级快速评估，{r['note']}。由 swmm-copilot 自动生成。")
    doc.save(str(path))
